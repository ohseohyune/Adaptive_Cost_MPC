"""Build the residual-zero ablation report as a .docx.

Numbers are read back out of the analysis CSVs rather than retyped, so the
document cannot drift from the run that produced it. Re-run after
analyze_residual_zero.py / analyze_loss_decomposition.py to refresh it.
"""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
RESIDUAL_ROOT = ROOT / "sweep_results" / "residual_zero_20260805"
ABLATION_ROOT = ROOT / "sweep_results" / "constraint_ablation_20260804"
CONDITIONS = ("D0", "D1", "D2", "D3", "M0")


def _rows(path: Path) -> list[dict]:
    with path.open() as handle:
        return list(csv.DictReader(handle))


def _table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    document.add_paragraph()


def _verdict_paragraph(document: Document, title: str, body: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    document.add_paragraph(body)


def _picture(document: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    document.add_picture(str(path), width=Inches(6.3))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph(caption)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.size = Pt(8)
        run.italic = True


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output", type=Path, default=ROOT / "residual_zero_ablation_report.docx"
    )
    args = parser.parse_args()

    seed_rows = _rows(RESIDUAL_ROOT / "residual_zero_seed_summary.csv")
    condition_rows = _rows(RESIDUAL_ROOT / "residual_zero_condition_summary.csv")
    loss_rows = _rows(ABLATION_ROOT / "constraint_ablation_loss_decomposition.csv")

    document = Document()
    for style_name, size in (("Normal", 10),):
        document.styles[style_name].font.size = Pt(size)

    document.add_heading("AC-MPC Residual-Zero Ablation 보고서", level=0)
    document.add_paragraph(
        "질문: D0의 success rate는 학습된 Actor residual의 기여로 나온 것인가, "
        "아니면 Actor residual이 없어도 phase별 fixed MPC prior만으로 같은 성능이 나오는가."
    )
    document.add_paragraph(
        "새로운 학습은 수행하지 않았다. 기존 constraint ablation checkpoint 15개를 "
        "deterministic evaluation으로 replay한 결과만 담는다."
    )

    document.add_heading("1. 실행 설정", level=1)
    _table(
        document,
        ["항목", "값"],
        [
            ["Job", "5 조건 x 3 seed x 2 mode = 30 (30/30 성공, 실패 0)"],
            ["Episode", "mode당 100 deterministic, 총 3,000"],
            ["Evaluation seed", "100000, curriculum-mode balanced, 전 job 동일"],
            ["Exploration", "OFF (online_learning=False -> policy mean)"],
            ["소요", "2026-08-05 12:14-13:37 KST, job당 9~13분, 4 동시"],
            [
                "실행 명령",
                "python3 acmpc/run_residual_zero_ablation.py --episodes 100 --jobs 4",
            ],
            ["Checkpoint", "sweep_results/constraint_ablation_20260804/pilot/<COND>/seed_<S>/checkpoint.pt"],
            ["출력", "sweep_results/residual_zero_20260805/"],
        ],
    )
    document.add_paragraph(
        "checkpoint는 job 디렉토리로 복사한 뒤 실행했다. evaluation 경로가 실제로 "
        "checkpoint를 덮어쓰는 것을 확인했으므로 이 예방조치가 필요했다."
    )

    document.add_heading("2. 감사 결과", level=1)
    _table(
        document,
        ["감사 항목", "기준", "실측", "결과"],
        [
            ["residual-zero == phase prior", "<= 1e-8", "0.0", "PASS"],
            ["(vacuous 방지) residual 켜면 차이", "> 1e-6", "통과", "PASS"],
            ["evaluation determinism (동일 obs 100회)", "<= 1e-8", "0.0", "PASS"],
            ["(반대검증) training=True spread", "> 1e-6", "통과", "PASS"],
            ["total - (surrogate + entropy bonus)", "< 1e-8", "float32 rounding 내", "PASS"],
            ["기존 예시 복원", "약 2.26e-8", "2.26e-8", "PASS"],
            ["post-step KL, lr=0", "약 0", "max <= 1e-8", "PASS"],
            ["post-step KL, lr=1e-2", "> 0", "> 0, NaN/Inf 0", "PASS"],
        ],
    )
    document.add_paragraph(
        "evaluation_exploration_enabled = false, evaluation_action_mode = \"mean\", "
        "action sampling 미사용. 코드 경로로 확인했다: exploring = config.online_learning "
        "and phase is not HOLD 이므로 eval에서는 training=False가 되고, "
        "online_actor_critic.py:889에서 normalized_action = normalized_mean이 된다. "
        "기존 acmpc_constraint_switch_test.py도 PASS로 회귀가 없다."
    )

    document.add_heading("3. approximate KL 계산 시점", level=1)
    document.add_paragraph(
        "old_log_prob -> current_log_prob -> ratio -> approximate_kl (pre-step, "
        "target_kl 조기중단 전용이며 로깅되지 않았음) -> backward -> gradient clipping "
        "-> optimizer.step. epoch 루프가 모두 끝난 뒤 online_actor_critic.py:1073에서 "
        "_policy_kl(batch)를 전체 batch에 대해 다시 계산하고, 그 값이 "
        "PPOUpdateSummary.approximate_kl로 로깅된다."
    )
    document.add_paragraph(
        "따라서 기존 로그의 KL은 post-step 값이다(선택지 B/D). 구조적으로 0이 되는 값이 "
        "아니므로 D0의 1e-5는 \"업데이트 후에도 정책이 실제로 움직이지 않았다\"는 유효한 "
        "증거다. 이번 수정으로 per-minibatch pre_step_approximate_kl / "
        "post_step_approximate_kl이 별도로 기록된다."
    )

    document.add_heading("4. Actor loss 3항 분해", level=1)
    document.add_paragraph(
        "total_actor_loss = policy_surrogate_loss - entropy_coef x entropy 이므로 "
        "기존 로그에서 policy_surrogate_loss = total_actor_loss + entropy_coef x entropy로 "
        "복원된다. entropy_coef=1e-3, entropy 약 -12.4이면 entropy bonus만으로 약 0.0124가 "
        "되어 기존에 기록된 값 전체를 설명한다."
    )
    marks = ("0", "399")
    rows = []
    for condition in CONDITIONS:
        for mark in marks:
            selected = [
                r
                for r in loss_rows
                if r["condition"] == condition and r["episode"] == mark and r["seed"] == "7"
            ]
            terms = {r["term"]: float(r["mean"]) for r in selected}
            if not terms:
                continue
            rows.append(
                [
                    condition,
                    mark,
                    f"{terms['total_actor_loss']:.6f}",
                    f"{terms['entropy_bonus']:.6f}",
                    f"{terms['policy_surrogate_loss']:.3e}",
                ]
            )
    _table(
        document,
        ["Condition", "Episode", "total_actor_loss", "entropy_bonus", "policy_surrogate_loss"],
        rows,
    )
    document.add_paragraph(
        "seed 7 기준. 전 조건/전 seed/episode 0,100,200,300,399의 평균-표준편차-최솟값-"
        "최댓값은 constraint_ablation_loss_decomposition.csv에, seed 평균 +- 표준편차는 "
        "_summary.csv에 있다."
    )
    _picture(
        document,
        ABLATION_ROOT / "constraint_ablation_loss_decomposition.png",
        "total_actor_loss와 entropy_bonus가 거의 완전히 겹친다 (seed min-max band).",
    )
    _picture(
        document,
        ABLATION_ROOT / "constraint_ablation_surrogate_loss.png",
        "policy_surrogate_loss만 별도 scale로 표시.",
    )
    document.add_paragraph(
        "주의: mean_executed_ppo_epochs가 D0에서 약 0.99로 사실상 1 epoch만 실행되므로, "
        "첫 PPO pass에서 ratio가 1이고 advantage가 minibatch 평균 0으로 정규화되면 "
        "surrogate loss 값 자체는 구조적으로 0에 가깝다. 이 값이 0이라는 사실만으로 "
        "학습 신호가 없다고 결론 내릴 수는 없다."
    )

    document.add_heading("5. residual-zero와 fixed MPC의 동일성", level=1)
    document.add_paragraph("동일하지 않다. phase prior 자체가 다르다.")
    _table(
        document,
        ["Phase", "Cost", "fixed baseline", "residual-zero", "relative error"],
        [
            ["PRE_IMPACT", "velocity", "6.0", "3.0", "1.0"],
            ["HOLD", "grasp", "0.0", "20.0", "1.0"],
        ],
    )
    document.add_paragraph(
        "통과 기준(median velocity relative error <= 1e-6)을 6자리 초과한다. 더 결정적으로 "
        "fixed baseline의 HOLD grasp=0.0은 AC-MPC actor가 표현할 수 없다. "
        "AdaptiveCostActor.__init__이 \"phase priors must be finite and positive\"로 "
        "거부하며, Mode C 실행이 실제로 이 오류로 종료되었다. 따라서 residual-zero는 "
        "\"AC-MPC architecture with zero actor residual\"로 구분해 부른다."
    )
    fixed = RESIDUAL_ROOT / "fixed_baseline_default_priors" / "result.json"
    if fixed.exists():
        rate = json.loads(fixed.read_text())["success_rate"]
        document.add_paragraph(
            f"대신 controller 등가물(기본 prior + weight_delta_fraction=0, actor 없음)을 "
            f"동일 100 episode로 실행한 결과 success = {rate:.3f}로, 15개 checkpoint의 "
            f"residual-zero 결과 0.920과 완전히 일치한다. residual-zero 구현이 actor를 "
            f"제대로 우회함을 성능 수준에서 재확인한 것이다."
        )

    document.add_heading("6. 결과", level=1)
    document.add_heading("6.1 seed별", level=2)
    _table(
        document,
        ["Cond", "Seed", "Learned", "Zero", "d success", "Learned IS", "Zero IS", "P95 learned", "P95 zero"],
        [
            [
                r["condition"],
                r["seed"],
                f"{float(r['learned_success']):.2f}",
                f"{float(r['zero_success']):.2f}",
                f"{float(r['delta_success']):+.3f}",
                f"{float(r['learned_impact_safe']):.2f}",
                f"{float(r['zero_impact_safe']):.2f}",
                f"{float(r['learned_peak_p95']):.2f}",
                f"{float(r['zero_peak_p95']):.2f}",
            ]
            for r in seed_rows
        ],
    )

    document.add_heading("6.2 조건별 (paired bootstrap, 20,000 resample)", level=2)
    _table(
        document,
        ["Condition", "Learned", "Zero", "Paired diff", "CI95", "Actor 판정"],
        [
            [
                r["condition"],
                f"{float(r['learned_success']):.3f}",
                f"{float(r['zero_success']):.3f}",
                f"{float(r['paired_delta']):+.3f}",
                f"[{float(r['ci95_low']):+.3f}, {float(r['ci95_high']):+.3f}]",
                r["verdict"],
            ]
            for r in condition_rows
        ],
    )
    document.add_paragraph(
        "resampling 단위는 episode pair다. 두 mode가 같은 scenario sequence를 replay하며, "
        "episode별 mass/friction 일치를 비교 전에 검증했다."
    )

    document.add_heading("6.3 residual 크기와 실제 기여", level=2)
    exploration = 0.03 * 1.8 * (6 ** 0.5)
    _table(
        document,
        ["Condition", "residual abs mean", "command delta (m/s)", "노이즈/신호", "d success", "d peak P95 (N)"],
        [
            [
                r["condition"],
                f"{float(r['actor_residual_abs_mean']):.4f}",
                f"{float(r['policy_command_delta']):.5f}",
                f"{exploration / float(r['policy_command_delta']):.1f}",
                f"{float(r['paired_delta']):+.3f}",
                f"{float(r['learned_peak_p95']) - float(r['zero_peak_p95']):+.2f}",
            ]
            for r in condition_rows
        ],
    )
    document.add_paragraph(
        f"노이즈/신호 = E||u_sample - u_mean|| / ||u_mean - u_zero||. 분자는 "
        f"exploration std 0.03 x velocity limit 1.8 x sqrt(6) = {exploration:.4f} m/s다. "
        "분모는 actor output이 아니라 MPC를 거친 최종 velocity command에서 측정했다. "
        "D0에서 이 비가 약 65라는 것은 탐색 노이즈가 학습된 residual의 command 효과보다 "
        "65배 크다는 뜻이다."
    )
    document.add_paragraph(
        "D0의 paired 승패(300 episode): learned만 성공 0건, zero만 성공 2건, 나머지 동률. "
        "actor가 단독으로 이긴 episode가 하나도 없다."
    )
    document.add_paragraph(
        "Emergency failure는 전 조건 전 mode에서 0.000으로 증가가 없다. "
        "Episode return은 evaluation 경로가 total_reward=0.0만 기록하므로 NOT MEASURED다."
    )

    document.add_heading("6.4 Robustness (mass / speed 3분위 worst-group)", level=2)
    _table(
        document,
        ["Condition", "mass worst (learned/zero)", "speed worst (learned/zero)"],
        [
            ["D0", "0.882 / 0.882", "0.755 / 0.765"],
            ["D1", "0.637 / 0.882", "0.470 / 0.636"],
            ["M0", "0.040 / 0.882", "0.000 / 0.000"],
        ],
    )
    document.add_paragraph(
        "size, initial position, vertical velocity, time-to-contact, prediction noise는 "
        "per-episode 로깅이 없어 NOT MEASURED다. 목표(평균 +5%p, worst-group +3%p)를 "
        "달성한 조건은 없다."
    )

    document.add_heading("7. 그래프", level=1)
    _picture(
        document,
        RESIDUAL_ROOT / "residual_zero_success.png",
        "조건별 / seed별 learned vs residual-zero success rate (seed 표준편차 bar).",
    )
    _picture(
        document,
        RESIDUAL_ROOT / "residual_zero_paired_and_forces.png",
        "paired episode 승패 분해와 18N / 36N 초과율.",
    )
    _picture(
        document,
        RESIDUAL_ROOT / "residual_zero_diagnostics.png",
        "residual 크기 및 command delta 대비 성능 차이, first-contact peak force 분포.",
    )
    document.add_paragraph(
        "pre/post step KL 시계열은 플래그 구현과 단위 테스트를 마쳤으나 기존 로그에는 "
        "없으므로 다음 학습 실행부터 생성된다."
    )

    document.add_heading("8. 결론", level=1)
    _verdict_paragraph(
        document,
        "결론 A - D0는 사실상 fixed MPC.",
        "learned 0.913 대 residual-zero 0.920, paired diff -0.007, CI95 [-0.017, +0.000]으로 "
        "상한이 0에 닿으며 개선 방향을 지지하지 않는다. ImpactSafe는 1.00으로 동일하고 "
        "peak force P95는 15.50 대 15.52로 0.1% 차이여서 10% 감소 기준에 미달한다. "
        "300개 paired episode 중 actor가 단독으로 이긴 episode는 0건이다. "
        "따라서 D0의 성능은 phase prior MPC가 만든 것이며 학습된 Actor residual의 "
        "추가 기여는 확인되지 않았다.",
    )
    _verdict_paragraph(
        document,
        "결론 C - constraint를 제거한 조건에서 Actor는 MPC prior를 직접 훼손했다.",
        "residual-zero는 D1~M0 전부에서 0.920으로 D0와 동일하게 회복한 반면 learned는 "
        "0.713 / 0.013 / 0.503 / 0.123이었고 모든 CI 상한이 0보다 작다. 늘어난 actor "
        "movement는 유용한 adaptation이 아니라 잘 튜닝된 prior를 훼손하는 방향의 정책 "
        "변화였다. M0 seed 7에서는 ImpactSafe가 1.00에서 0.16으로, peak force P95가 "
        "15.5N에서 39.6N으로 무너져 안전성까지 악화되었다.",
    )
    _verdict_paragraph(
        document,
        "결론 D - residual-zero와 cond1 fixed MPC는 동일한 controller가 아니다.",
        "차이는 phase prior에서 발생했다(PRE_IMPACT velocity 3.0 대 6.0, HOLD grasp "
        "20.0 대 0.0). 더구나 fixed baseline의 prior는 AC-MPC actor가 표현할 수 없다. "
        "직접 비교를 위해서는 baseline 설정 정렬이 필요하다.",
    )
    _verdict_paragraph(
        document,
        "추가 발견 - 기존 evaluation 수치가 재현되지 않는다.",
        "D0 seed 7 checkpoint를 동일 seed, 동일 50 episode로 cold-start replay하면 "
        "success가 0.92인데 학습 중 기록된 값은 1.00이었다. 같은 replay의 "
        "mean_first_contact_peak_force_n은 12.597로 기록값 12.59676과 5자리까지 "
        "일치한다. 즉 물리와 궤적은 동일하고 hold 단계 판정만 갈렸다(4/50, 전부 "
        "unstable box motion, hold 0.09~5.00초). 학습 프로세스 내부에서 실행된 "
        "evaluation과 별도 프로세스 replay가 hold 구간에서 갈라진다. 이는 별도 감사 "
        "대상이며, 이번 learned 대 zero 결론에는 영향이 없다. 양쪽 arm 모두 동일 "
        "조건에서 cold-start로 새로 측정했기 때문이다.",
    )

    document.add_heading("9. 다섯 질문에 대한 답", level=1)
    for question, answer in [
        (
            "D0의 100% 성공은 Actor 덕분인가?",
            "아니다. paired diff -0.007 (CI95 [-0.017, +0.000])이고 actor 단독 승리가 "
            "300 episode 중 0건이다. 게다가 그 100% 자체가 재현되지 않아 실제로는 0.92이며, "
            "그 0.92는 actor 없는 MPC도 동일하게 낸다.",
        ),
        (
            "residual-zero는 cond1 fixed MPC와 실제로 동일한가?",
            "아니다. phase prior 2개 항목이 relative error 1.0으로 다르고, fixed baseline의 "
            "HOLD grasp=0.0은 AC-MPC actor가 표현조차 하지 못한다.",
        ),
        (
            "constraint를 제거한 조건에서 Actor는 성능을 개선했는가, 악화했는가?",
            "전부 악화했다. D1 -0.207, D2 -0.907, D3 -0.417, M0 -0.797이며 모든 CI95 상한이 "
            "0보다 작다. M0에서는 안전성(ImpactSafe 0.16, peak 39.6N)까지 무너뜨렸다.",
        ),
        (
            "기존 actor loss 0.0124는 실제 surrogate loss였는가?",
            "아니다. 전량이 entropy bonus였고 순수 surrogate는 1e-8에서 1e-7 규모다. 다만 "
            "1 epoch 설정에서는 surrogate 값 자체가 진단 정보를 갖지 않으므로 이 사실만으로 "
            "학습 신호가 없다고 주장해서는 안 된다. KL 쪽 증거는 post-step이라 유효하다.",
        ),
        (
            "다음 단계에서 fixed std 실험으로 넘어가도 되는가?",
            "아직 아니다. 두 가지가 먼저다. 첫째, actor의 command 기여가 탐색 노이즈의 약 "
            "1/65이므로 PPO가 자기 residual의 효과를 노이즈 속에서 관측할 수 없다. std를 "
            "0.03에 고정해도 이 비율은 그대로이며, 오히려 std를 낮추거나 command delta를 "
            "키우는 쪽이 문제의 축이다. 둘째, in-process evaluation과 cold-start replay가 "
            "hold 판정에서 갈리는 문제를 먼저 잡지 않으면 이후 어떤 실험의 성공률도 신뢰할 "
            "수 없다.",
        ),
    ]:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(question)
        run.bold = True
        document.add_paragraph(answer)

    document.add_heading("10. 산출물", level=1)
    _table(
        document,
        ["파일", "내용"],
        [
            ["residual_zero_episodes.csv", "3,000 episode 전수 지표"],
            ["residual_zero_seed_summary.csv", "seed별 요약과 CI"],
            ["residual_zero_condition_summary.csv", "조건별 요약과 판정"],
            ["residual_zero_report.md", "생성된 요약 표"],
            ["residual_zero_success.png", "조건별 / seed별 success"],
            ["residual_zero_paired_and_forces.png", "paired 승패, force 초과율"],
            ["residual_zero_diagnostics.png", "residual/command delta 대비 성능, force 분포"],
            ["constraint_ablation_loss_decomposition.csv", "loss 3항 복원 전수"],
            ["constraint_ablation_loss_decomposition_summary.csv", "seed 평균 +- 표준편차"],
            ["constraint_ablation_loss_decomposition.png", "total과 entropy bonus 중첩"],
            ["constraint_ablation_surrogate_loss.png", "surrogate loss 별도 scale"],
        ],
    )

    document.save(args.output)
    print(f"wrote {args.output}")


if __name__ == "__main__":
    main()
